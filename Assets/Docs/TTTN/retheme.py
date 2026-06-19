# -*- coding: utf-8 -*-
"""Re-theme GameDesignDocument_BuonVillage (Tay Nguyen) -> LANG BANH (banh truyen thong VN).
Giu nguyen dinh dang/bang/font; chi thay lop van hoa. Bam sat GDD goc."""
import io, sys
from docx import Document

SRC = r"D:/Ptiter/TTTN/GameDesignDocument_BuonVillage.docx"
DST = r"D:/Ptiter/TTTN/GameDesignDocument_LangBanh.docx"

# Thu tu: cum dai/dac trung truoc, fallback chung sau.
REPLACEMENTS = [
    # ---- Tieu de / nhan dien ----
    ("BUÔN VILLAGE", "LÀNG BÁNH"),
    ("Buôn Village — Game Design Document", "Làng Bánh — Game Design Document"),
    ("Game Farming / Life Simulation 2.5D — Mobile-first — Bản sắc Tây Nguyên",
     "Game Farming / Cooking Simulation 2.5D — Mobile-first — Bản sắc ẩm thực Việt"),

    # ---- 1. TOM TAT ----
    ("Buôn Village — game Farming / Life Sim mobile-first, bối cảnh buôn làng Tây Nguyên, kết hợp vòng lặp gây nghiện kiểu Dreamdale với chất nghệ thuật vẽ tay sử thi kiểu Don't Starve.",
     "Làng Bánh — game Farming / Cooking Sim mobile-first, bối cảnh làng quê Việt với nghề làm bánh truyền thống, kết hợp vòng lặp gây nghiện kiểu Dreamdale với chất nghệ thuật vẽ tay mộc mạc kiểu Don't Starve."),
    ("Chưa game farming nào khai thác văn hoá Tây Nguyên → khác biệt rõ trên thị trường.",
     "Chưa game farming nào khai thác văn hoá bánh truyền thống Việt Nam → khác biệt rõ trên thị trường."),

    # ---- 2. TONG QUAN ----
    ("Farming + Life Simulation + Adventure RPG nhẹ",
     "Farming + Cooking Simulation + Adventure RPG nhẹ"),
    ("Bản sắc Tây Nguyên · cơ chế một-chạm gây nghiện · âm nhạc cồng chiêng",
     "Bản sắc ẩm thực bánh Việt · cơ chế một-chạm gây nghiện · âm nhạc dân gian (sáo trúc, đàn bầu)"),

    # ---- 3. CORE LOOP (so do ASCII) ----
    ("[Di chuyen + Thu thap mot-cham: Go / Da / Ca phe]",
     "[Di chuyen + Thu thap mot-cham: Cui / La goi / Nep]"),
    ("[Ban tho HOAC Che bien sau -> Dac san (gia X3-X5)]",
     "[Ban tho HOAC Che bien -> Banh dac san (gia X3-X5)]"),
    ("[Tich luy Tien te + Danh tieng Buon lang]",
     "[Tich luy Tien te + Danh tieng Lang nghe]"),
    ("[Nang cap Nha Dai + Mo rong Nong trai]",
     "[Nang cap Lo Banh + Mo rong Nong trai]"),

    # ---- 4.1 Dieu khien ----
    ("cạnh cây → tự vung rìu; cạnh cà phê chín → tự tuốt vào gùi.",
     "cạnh cây → tự vung rìu lấy củi; cạnh lúa nếp chín → tự gặt vào thúng."),
    ("Giới hạn gùi: Sức chứa gùi có hạn → tạo nhịp di chuyển có chủ đích.",
     "Giới hạn quang gánh: Sức chứa quang gánh có hạn → tạo nhịp di chuyển có chủ đích."),

    # ---- 4.2 Land Expansion ----
    ("MVP chỉ làm: khu Nông trại gia đình (nhà dài cấp 1, giếng cổ, ô rẫy bazan).",
     "MVP chỉ làm: khu Nông trại gia đình (bếp/lò bánh cấp 1, giếng làng, luống ruộng phù sa)."),
    ("phát cutscene mở ra Buôn làng như lời hứa hẹn cho người chơi.",
     "phát cutscene mở ra Làng nghề / Hội làng bánh như lời hứa hẹn cho người chơi."),
    ("mở rộng đầy đủ các phân khu (Buôn làng/Hub, Rừng Già, Bến Nước)",
     "mở rộng đầy đủ các phân khu (Làng nghề/Hub, Chợ Phiên, Bến Nước)"),

    # ---- 4.3 Farming (bang) ----
    ("Lúa Rẫy", "Lúa Nếp"),
    ("Hạt lúa thô", "Gạo nếp thô"),
    ("Cà Phê Robusta", "Mía"),
    ("Quả cà phê chín", "Cây mía (mật/đường)"),
    ("Hồ Tiêu", "Đậu Xanh"),
    ("Chuỗi tiêu xanh", "Chùm đậu xanh"),
    ("Cần cọc gỗ để leo (5 gỗ)", "Cần làm giàn leo (5 gỗ)"),

    # ---- 4.4 Chan nuoi (bang) ----
    ("Gà Rừng", "Gà Ri"),
    ("Ăn lúa thô → đẻ Trứng mỗi 45s; đi ngang để nhặt",
     "Ăn thóc/nếp → đẻ Trứng mỗi 45s; đi ngang để nhặt"),
    ("Heo Sọc Dưa", "Lợn Ỉ"),
    ("Ăn Sắn/Khoai → Thịt Heo Gác Bếp (qua lò sấy)",
     "Ăn Sắn/Khoai → Thịt làm nhân bánh (qua bếp lò)"),
    ("Trâu Bản Địa", "Trâu Cày"),
    ("+50% tốc độ & sức chứa gùi trong nông trại (buff cơ động)",
     "+50% tốc độ & sức chứa quang gánh trong nông trại (buff cơ động)"),

    # ---- 4.5 Crafting pipeline ----
    ("Ca Phe Tho -> [San Phoi] -> Nhan Xanh -> [Lo Rang Dat Set] -> Ca Phe Rang Cui  (X5)",
     "Gao Nep -> [Coi Gia Bot] -> Bot Nep -> [Goi La Dong + Noi Dong Luoc] -> Banh Chung  (X5)"),
    ("Tieu Xanh  -> [Gui Say Kho] -> Ho Tieu Den Bao Tho Cam                          (X3)",
     "Dau Xanh -> [Cho Do Hap] -> Nhan Dau + Mat Mia -> [Khuon Go] -> Banh Dau Xanh   (X3)"),

    # ---- 4.6 Progression (bang) ----
    ("Cấp Nhà Dài", "Cấp Lò Bánh"),
    ("Cấp 1 — Nhà Sàn Nhỏ", "Cấp 1 — Bếp Tranh Nhỏ"),
    ("Cấp 2 — Mở Rộng", "Cấp 2 — Mở Rộng Lò"),
    ("NPC Thương lái; +4 ô rẫy", "NPC Thương lái; +4 thửa ruộng"),
    ("Cấp 3 — Đại Bản Doanh", "Cấp 3 — Hiệu Bánh Làng Nghề"),
    ("300 Gỗ Quý + 150 Cà phê Rang", "300 Gỗ Quý + 150 Bánh Chưng"),
    ("Kích hoạt cutscene mở bản làng (kết mở)",
     "Kích hoạt cutscene mở làng nghề (kết mở)"),

    # ---- 4.7 NPC ----
    ("Già Làng Ama Thuột: Main Quest — yêu cầu tài nguyên sửa công trình.",
     "Già Làng / Cụ Đồ làng: Main Quest — yêu cầu tài nguyên sửa đình làng / lò bánh tổ."),
    ("Nghệ nhân Y-Khuê: Side Quest — cung cấp nông sản chuẩn → bản vẽ nâng cấp.",
     "Nghệ nhân Bà Hương: Side Quest — cung cấp nông sản chuẩn → công thức bánh nâng cấp."),
    ("Thương Lái Phương Xa: thu mua nông sản chế biến, giá biến động theo ngày.",
     "Thương Lái Phương Xa: thu mua bánh đặc sản, giá biến động theo phiên chợ."),

    # ---- 4.8 UI/UX & Audio ----
    ("UI: triển khai theo UICollection (mục 5.4); panel vân gỗ-dây thừng; thanh trạng thái dạng gùi đầy nước/củi cháy; icon vẽ tay thổ cẩm.",
     "UI: triển khai theo UICollection (mục 5.4); panel vân gỗ-tre-lạt; thanh trạng thái dạng quang gánh / nồi bánh đang luộc; icon vẽ tay hoa văn trống đồng - giấy dó."),
    ("Audio: ve sầu/gió tre/thác nước; SFX \"cộc cộc\", \"xoàn xoạt\"; BGM đàn T'rưng + cồng chiêng + trống da trâu.",
     "Audio: gà gáy/gió đồng/lửa lò bập bùng; SFX \"cộc cộc\" (giã bột), \"lóc bóc\" (nồi luộc); BGM sáo trúc + đàn bầu + trống chèo."),

    # ---- 5.6 Luong du lieu (ASCII) ----
    ("ban event UI cap nhat gui", "ban event UI cap nhat quang ganh"),

    # ---- 6. MVP ----
    ("Chế biến cà phê + Progression Nhà Dài (cấp 1→3).",
     "Chế biến bánh + Progression Lò Bánh (cấp 1→3)."),
    ("Cutscene mở bản làng (kết mở) — thay cho mini-game.",
     "Cutscene mở làng nghề (kết mở) — thay cho mini-game."),
    ("Mini-game Lễ hội Cồng chiêng (rhythm) — đẩy sang tương lai.",
     "Mini-game Hội thi gói bánh (rhythm) — đẩy sang tương lai."),
    ("Mở rộng đầy đủ bản đồ (Hub, Rừng Già, Bến Nước), câu cá, bẫy tôm.",
     "Mở rộng đầy đủ bản đồ (Hub, Chợ Phiên, Bến Nước), câu cá, bẫy tôm."),

    # ---- 7.1 Doi tuong & Thi truong ----
    ("Lợi thế: ngách văn hoá Tây Nguyên chưa đối thủ chiếm.",
     "Lợi thế: ngách văn hoá ẩm thực bánh truyền thống Việt chưa đối thủ chiếm."),

    # ---- 7.2 Sprint (bang) ----
    ("Di chuyển + context-action + thu thập + gùi + bán cơ bản; hoàn thiện hệ thống Map nông trại (grid/đặt object)",
     "Di chuyển + context-action + thu thập + quang gánh + bán cơ bản; hoàn thiện hệ thống Map nông trại (grid/đặt object)"),
    ("Pipeline cà phê; Nhà Dài 1→3; NPC quest + bán; Save/Load Firebase",
     "Pipeline bánh; Lò Bánh 1→3; NPC quest + bán; Save/Load Firebase"),
    ("Cutscene mở bản làng; UI theo UICollection; audio, hiệu ứng, fix bug; build demo",
     "Cutscene mở làng nghề; UI theo UICollection; audio, hiệu ứng, fix bug; build demo"),

    # ---- 7.3 KPI ----
    ("Nâng cấp Nhà Dài tới cấp 3 → kích hoạt cutscene mở bản làng.",
     "Nâng cấp Lò Bánh tới cấp 3 → kích hoạt cutscene mở làng nghề."),

    # ---- 8. POST-MVP ----
    ("Mini-game Lễ hội Cồng chiêng (rhythm) + đa lễ hội (đâm trâu, mừng lúa mới).",
     "Mini-game Hội thi gói bánh (rhythm) + đa lễ hội (Tết, mừng lúa mới, hội làng)."),
    ("Mở rộng đầy đủ bản đồ: Hub, Rừng Già, Bến Nước; câu cá, bẫy tôm.",
     "Mở rộng đầy đủ bản đồ: Hub, Chợ Phiên, Bến Nước; câu cá, bẫy tôm."),
    ("Kinh tế động (Chợ Buôn giá cung–cầu); thời tiết, mùa vụ.",
     "Kinh tế động (Chợ Phiên giá cung–cầu); thời tiết, mùa vụ."),
    ("Backend riêng: authoritative time, anti-cheat, bảng xếp hạng buôn làng.",
     "Backend riêng: authoritative time, anti-cheat, bảng xếp hạng làng nghề."),
    ("Thương mại hoá: cosmetic thổ cẩm, không pay-to-win.",
     "Thương mại hoá: cosmetic gói bánh / giấy dó, không pay-to-win."),

    ("tiếp cận ô rẫy → tự gieo", "tiếp cận ô ruộng → tự gieo"),

    # ---- Fallback chung (sau cung) ----
    ("Tây Nguyên", "làng quê Việt"),
    ("buôn làng", "làng nghề"),
]

def replace_in_paragraph(p, old, new):
    """Thay the across-runs, giu format run dau cua match."""
    runs = p.runs
    if not runs:
        return 0
    count = 0
    search_from = 0
    while True:
        full = "".join(r.text for r in runs)
        idx = full.find(old, search_from)
        if idx == -1:
            break
        end = idx + len(old)
        search_from = idx + len(new)  # skip qua phan vua chen -> tranh loop vo han
        pos = 0
        inserted = False
        for r in runs:
            rstart, rend = pos, pos + len(r.text)
            pos = rend
            ov_s, ov_e = max(idx, rstart), min(end, rend)
            if ov_s < ov_e:
                before = r.text[:ov_s - rstart]
                after = r.text[ov_e - rstart:]
                if not inserted:
                    r.text = before + new + after
                    inserted = True
                else:
                    r.text = before + after
        count += 1
    return count

def iter_paragraphs(container):
    for p in container.paragraphs:
        yield p
    for t in container.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)

def main():
    doc = Document(SRC)
    paras = list(iter_paragraphs(doc))
    for sec in doc.sections:
        paras += list(iter_paragraphs(sec.header))
        paras += list(iter_paragraphs(sec.footer))
    stats = {}
    for old, new in REPLACEMENTS:
        c = 0
        for p in paras:
            c += replace_in_paragraph(p, old, new)
        stats[old] = c
    doc.save(DST)
    miss = [o for o, c in stats.items() if c == 0]
    print("Saved:", DST)
    print("Tong thay the:", sum(stats.values()))
    print("KHONG khop (can kiem tra):")
    for m in miss:
        print("  -", m[:70])

if __name__ == "__main__":
    main()
