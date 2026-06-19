# -*- coding: utf-8 -*-
"""Sinh GameDesignDocument_BepViet.docx — chu de Do an Viet Nam.
Bam sat checklist GDD / Developer / Art do user cung cap.
Giu trang bia + nhan su goc cua Buon Village GDD."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DST = r"D:/Ptiter/TTTN/GDD_ThucTapTotNghiep.docx"
ACCENT = RGBColor(0xB1, 0x3A, 0x1E)   # do gach - sac do an Viet
DARK   = RGBColor(0x33, 0x26, 0x1D)

doc = Document()

# ---------- base style ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

def style_heading(level, color):
    pass

for hn, sz, col in [("Heading 1", 15, ACCENT), ("Heading 2", 12.5, DARK)]:
    st = doc.styles[hn]
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.name = "Calibri"
    st.font.bold = True

def h1(t):
    p = doc.add_heading(t, level=1)
    return p

def h2(t):
    return doc.add_heading(t, level=2)

def para(t="", bold=False, italic=False, size=11, color=None, align=None, space_after=4):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p, r

def lead(label, rest):
    """Bullet voi phan dau in dam."""
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(label); r.bold = True
    p.add_run(rest)
    return p

def bullet(t):
    return doc.add_paragraph(t, style="List Bullet")

def mono(lines):
    for ln in lines:
        p = doc.add_paragraph()
        r = p.add_run(ln)
        r.font.name = "Consolas"; r.font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        rp = hdr[i].paragraphs[0].add_run(htext)
        rp.bold = True; rp.font.size = Pt(10.5); rp.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "B13A1E")
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(str(val))
            rr.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def caption(t):
    p, r = para(t, italic=True, size=9.5, color=RGBColor(0x70,0x70,0x70),
                align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    return p

def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    t = OxmlElement("w:t"); t.text = "Nhan chuot phai → Update Field de hien Muc luc."
    r = OxmlElement("w:r"); r.append(t); fld.append(r)
    p._p.append(fld)

# =================================================================
# TRANG BIA
# =================================================================
para("HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG", bold=True, size=13,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Cơ sở: Thành Phố Hồ Chí Minh", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para("Khoa: Công Nghệ Thông Tin 2", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("━━━━━━━━━━━━━━━", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
para("[ TÊN GAME — ĐANG CẬP NHẬT ]", bold=True, size=26, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("Game Farming + Cooking Simulation — Mobile (Landscape) · 60FPS · Offline/Online",
     italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Bản sắc Đồ ăn Việt Nam — Từ nông trại đến mâm cơm Việt",
     italic=True, size=12, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
para("BẢN THẢO ĐỀ XUẤT DỰ ÁN THỰC TẬP", bold=True, size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

para("Tác giả:", bold=True, space_after=0)
para(" Vũ Đình Khoa — PO / Lead Dev / Game Designer", space_after=6)
para("Thành viên CNTT:", bold=True, space_after=0)
para(" Vũ Đình Khoa — N22DCCN142 — D22CQCNPM02-N", space_after=0)
para(" Phạm Hùng Thiên — N22DCCN180 — D22CQCNPM02-N", space_after=0)
para(" Văn Minh Tấn — N22DCCN175 — D22CQCNPM02-N", space_after=6)
para("Thành viên khác:", bold=True, space_after=0)
para(" Đội Artist ngành Đa phương tiện", space_after=10)
para("Tháng 06 năm 2026", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

doc.add_page_break()

# =================================================================
# GHI CHU NHAN SU + MUC LUC
# =================================================================
h1("GHI CHÚ NHÂN SỰ DỰ ÁN")
lead("Leader: ", "Vũ Đình Khoa — Product Owner, Lead Developer, Game Designer.")
lead("Phạm Hùng Thiên: ", "Developer, Game Design.")
lead("Văn Minh Tấn: ", "Developer.")
lead("Mảng Art: ", "03 sinh viên ngành Đa phương tiện đảm nhận.")

h1("MỤC LỤC")
add_toc()
doc.add_page_break()

# =================================================================
# 1. TOM TAT
# =================================================================
h1("1. TÓM TẮT")
para("Dự án — game Farming + Cooking Simulation mobile (màn hình ngang, 60FPS, "
     "chơi được offline lẫn online). Người chơi trồng trọt – chăn nuôi để lấy nguyên liệu, "
     "rồi chế biến thành các món ăn Việt Nam quen thuộc (bánh mì, phở, cơm tấm, bánh chưng…). "
     "Mỗi món hoàn thành lần đầu mở ra một cutscene kể chuyện gắn với nhân vật nổi tiếng.")
lead("Bản sắc độc bản: ", "Chưa game farming nào hệ thống hoá ẩm thực Việt Nam (nguyên liệu → "
     "món ăn → câu chuyện) → khác biệt rõ và quảng bá văn hoá Việt.")
lead("Tâm lý cốt lõi: ", "Liên tục mở khoá nguyên liệu & món ăn MỚI để đánh vào tâm lý thích "
     "khám phá; checklist nguyên liệu + quest tạo mục tiêu ngắn hạn rõ ràng.")
lead("Khả thi kỹ thuật: ", "Đã có 04 module nền (Service + DI VContainer + ScriptableObject) "
     "đang vận hành; Map dựng trên Unity Tilemap; Time chống cheat qua Firebase → MVP khả thi 8 tuần.")
lead("Mục tiêu thực tập: ", "Demo dọc (vertical slice) chạy trọn core loop trên mobile landscape, "
     "đủ trình diễn & bảo vệ, publish CHPlay nếu có tiềm năng.")

# =================================================================
# 2. TONG QUAN
# =================================================================
h1("2. TỔNG QUAN DỰ ÁN")
table(["Thuộc tính", "Giá trị"], [
    ["Thể loại", "Farming + Cooking Simulation + Light RPG"],
    ["Reference", "Dreamdale (one-tap, land expansion) · Cooking/Diner games · Stardew (farm→craft)"],
    ["Engine nền", "Unity (URP)"],
    ["Hướng màn hình", "Ngang (Landscape), khoá ngang — tối ưu 60 FPS"],
    ["Chế độ chơi", "Offline (chơi không mạng) + Online (cloud save, event, chống cheat thời gian)"],
    ["Phong cách hình ảnh", "2.5D / 2D Tilemap, vẽ tay, ấm áp ẩm thực"],
    ["Nền tảng", "Mobile — Android trước, iOS sau"],
    ["USP", "Đồ ăn Việt Nam · unlock liên tục nguyên liệu/món mới · cutscene gắn nhân vật nổi tiếng"],
], widths=[1.7, 4.6])

# =================================================================
# 3. CORE LOOP
# =================================================================
h1("3. VÒNG LẶP GAMEPLAY CỐT LÕI (CORE LOOP)")
mono([
 "[Chon mon muon lam]  (tu Unlock / Quest)",
 "        |",
 "        v",
 "[Xem checklist nguyen lieu can thiet]",
 "        |",
 "        v",
 "[Trong trot / Chan nuoi]  --theo Time System-->  [Thu hoach]",
 "        |",
 "        v",
 "[San xuat / Nau an]  (Cooking + Animation)",
 "        |",
 "        v",
 "[Hoan thanh mon]  -->  [CutScene cau chuyen mon an]",
 "        |",
 "        v",
 "[Coins + Danh tieng]  -->  [Unlock nguyen lieu / mon MOI]  (lap lai)",
])
caption("Sơ đồ 1 — Vòng lặp gameplay cốt lõi")
lead("Triết lý: ", "vòng lặp ngắn, phần thưởng tức thời (coins + món mới + cutscene); "
     "luôn hé lộ nội dung kế tiếp để duy trì động lực 'thích cái mới'.")

# =================================================================
# 4. THIET KE HE THONG GAMEPLAY
# =================================================================
h1("4. THIẾT KẾ HỆ THỐNG GAMEPLAY")

h2("4.1. Điều khiển & Camera — Mobile Landscape")
lead("One-tap context-aware: ", "chạm để di chuyển & tự thao tác (gần cây → thu hoạch, gần "
     "bếp → nấu); không cần đổi công cụ thủ công.")
lead("Drag bản đồ: ", "kéo (pan) để di chuyển camera quanh map; pinch-zoom (tuỳ chọn).")
lead("Khoá ngang: ", "UI bố trí ngang, mục tiêu 60 FPS ổn định trên thiết bị tầm trung.")
lead("PC (phụ): ", "WASD + chuột — chỉ phục vụ kiểm thử & phát triển.")

h2("4.2. Map System — 1 map, Unity Tilemap")
lead("Một map duy nhất: ", "tổ chức theo Tilemap của Unity (Grid + Tilemap layers).")
lead("Config được: ", "kích thước, vùng trồng, vật trang trí cấu hình qua ScriptableObject / "
     "Tilemap data — đổi map không cần sửa code.")
lead("Trồng lên map: ", "mỗi ô tile hợp lệ là 1 slot trồng cây/đặt chuồng.")
lead("Kéo (drag) map: ", "camera pan mượt trong biên map; có giới hạn (clamp) tránh lố biên.")

h2("4.3. Hệ thống Unlock")
para("Hai nhánh mở khoá song song, đều dùng Coins:")
lead("Unlock Nguyên liệu: ", "mở giống cây trồng / loài vật nuôi mới để sản xuất nguyên liệu.")
lead("Unlock Món ăn (Recipe): ", "mỗi món có checklist nguyên liệu; phải đủ nguyên liệu đã "
     "unlock + sản xuất được thì mới làm được món.")
lead("Cơ chế 'a Coins': ", "mỗi mục unlock có giá coins riêng; món/nguyên liệu kế tiếp luôn "
     "hiển thị dạng 'sắp mở' (teaser) để kích thích khám phá.")
para("Ví dụ checklist món (thiết kế trong ScriptableObject):", italic=True, space_after=2)
table(["Món ăn Việt", "Checklist nguyên liệu", "Unlock (coins)", "Giá bán"], [
    ["Bánh Mì (tutorial)", "Bột + Thịt + Rau + Trứng", "Miễn phí", "30"],
    ["Cơm Tấm", "Gạo + Thịt heo + Trứng + Rau", "300", "80"],
    ["Phở Bò", "Bánh phở (Gạo→Bột) + Thịt bò + Hành + Rau", "600", "150"],
    ["Gỏi Cuốn", "Bánh tráng + Tôm + Thịt + Rau", "800", "180"],
    ["Bún Chả", "Bún + Thịt heo + Rau + Nước mắm", "1000", "220"],
    ["Bánh Chưng (Event Tết)", "Gạo nếp + Đậu xanh + Thịt heo + Lá dong", "Mở qua Event", "300"],
], widths=[1.7, 2.9, 1.1, 0.7])

h2("4.4. Trồng trọt & Chăn nuôi (Farm System)")
para("Vận hành theo Time System; mỗi nguyên liệu có hình thái sinh trưởng theo giai đoạn "
     "(giống → đang phát triển → trưởng thành) để người chơi cảm nhận tiến trình.")
table(["Cây trồng", "Thời gian chín", "Giai đoạn", "Sản phẩm"], [
    ["Rau thơm", "30 giây", "3 giai đoạn", "Rau (gia vị/ăn kèm)"],
    ["Lúa", "60 giây", "3 giai đoạn", "Gạo (→ bột, bánh phở)"],
    ["Lúa Nếp", "90 giây", "3 giai đoạn", "Gạo nếp (bánh chưng)"],
    ["Đậu Xanh", "120 giây", "3 giai đoạn", "Đậu xanh (nhân)"],
    ["Mía", "180 giây (thu nhiều đợt)", "3 giai đoạn", "Đường / mật"],
], widths=[1.4, 2.0, 1.4, 1.5])
table(["Vật nuôi", "Cơ chế (theo Time System)", "Sản phẩm"], [
    ["Gà", "Ăn thóc → đẻ Trứng mỗi 45s; đi ngang để nhặt", "Trứng / Thịt gà"],
    ["Heo", "Ăn rau/sắn → tới chu kỳ thu", "Thịt heo (nhân/món)"],
    ["Bò", "Nuôi dài hơn → giá trị cao", "Thịt bò (phở)"],
], widths=[1.3, 3.5, 1.5])

h2("4.5. Sản xuất / Nấu ăn (Cooking System)")
lead("Trạm chế biến: ", "đưa nguyên liệu vào trạm (cối/bếp/nồi) → ra bán thành phẩm hoặc món.")
lead("Animation: ", "mỗi bước nấu có animation (khuấy, chiên, hấp) tạo cảm giác 'đang nấu'.")
lead("Hàng đợi: ", "có thể xếp nhiều món; thời gian nấu theo Time System (có thể tăng tốc qua Booster).")
para("Pipeline ví dụ:", italic=True, space_after=2)
mono([
 "Gao   -> [Coi/Xay] -> Bot -> [Lo Nuong/Hap] -> Banh pho / Vo banh",
 "Banh pho + Thit bo + Hanh + Nuoc dung -> [Bep Pho] -> PHO BO  (mon hoan chinh)",
])

h2("4.6. Hệ thống Quest")
lead("Ép trải nghiệm món mới: ", "ngay sau khi unlock 1 món → phát quest yêu cầu làm món đó "
     "N lần (đảm bảo người chơi thử nội dung vừa mở).")
lead("Phân loại: ", "Tutorial quest (dẫn dắt) · Main quest (theo tiến trình unlock) · "
     "Daily quest (giữ chân).")
lead("Phần thưởng: ", "coins, nguyên liệu hiếm, mở event / công thức tiếp theo.")

h2("4.7. Hệ thống Event")
lead("Vòng đời chung (LifeCycle): ", "Schedule → Activate → Progress/Milestone → Reward → Cleanup; "
     "dùng chung cho mọi event.")
lead("Config không cần code: ", "thời gian, món của event, phần thưởng, theme art… khai báo trong "
     "ScriptableObject → designer tự cấu hình.")
lead("Event đầu tiên — 'Tết Việt: Gói Bánh Chưng': ", "quảng bá văn hoá; mở công thức bánh chưng, "
     "trang trí map theo Tết, phần thưởng giới hạn thời gian.")
lead("Mở rộng tương lai: ", "Trung Thu (bánh trung thu), Giỗ Tổ, mùa lễ hội ẩm thực vùng miền.")

h2("4.8. Hệ thống Tutorial")
lead("Món thân thuộc — Bánh Mì Việt Nam: ", "chọn món ai cũng biết để giảm rào cản.")
lead("Từng bước: ", "di chuyển → trồng/thu nguyên liệu → nấu → hoàn thành bánh mì → nhận thưởng; "
     "có highlight, tooltip, hand-holding nhẹ.")
lead("Mục tiêu: ", "người chơi nắm trọn core loop trong < 3 phút đầu.")

h2("4.9. Shop & Booster")
lead("Booster tăng tốc: ", "x2 tốc độ trồng trọt · x2 tài nguyên thu hoạch · giảm thời gian nấu · "
     "mở rộng kho.")
lead("Tiền tệ: ", "mua bằng Coins (soft currency) trong MVP; IAP để hậu MVP.")

h2("4.10. Cutscene sáng tạo")
lead("Cơ chế: ", "hoàn thành lần đầu một món 'signature' → mở cutscene kể câu chuyện món ăn "
     "gắn với một nhân vật nổi tiếng (thành phẩm → câu chuyện).")
lead("Ví dụ ý tưởng: ", "Phở → nghệ nhân gánh phở Hà Nội xưa; Bánh mì → câu chuyện giao thoa "
     "ẩm thực Việt–Pháp.")
lead("Lưu ý pháp lý: ", "dùng nhân vật gốc do team sáng tạo hoặc đã được cấp phép để tránh vi "
     "phạm hình ảnh/bản quyền.")
lead("Mục tiêu: ", "phần thưởng cảm xúc + quảng bá văn hoá ẩm thực Việt.")

h2("4.11. UI/UX & Audio")
lead("UI: ", "triển khai theo UICollection (mục 5.7); bố cục ngang; panel ấm (gỗ/tre); "
     "checklist nguyên liệu trực quan; icon món vẽ tay.")
lead("Audio: ", "tiếng chợ/bếp/lửa; SFX nấu (xèo, lóc bóc, cộc cộc); BGM dân gian nhẹ "
     "(sáo trúc, đàn bầu).")

h2("4.12. Hệ thống Nâng cấp Công trình (Update System)")
para("Các công trình trong farm có thể xây & nâng cấp bằng Coins + nguyên liệu để tăng năng "
     "lực sản xuất — tạo trục tiến trình (progression) dài hạn, gắn coins kiếm được từ core loop.")
lead("Cơ chế: ", "chọn công trình → xem checklist yêu cầu (coins + nguyên liệu) → nâng cấp → "
     "công trình đổi hình theo cấp + nhận buff.")
lead("Đặt trên map: ", "công trình nằm trên Tilemap (mục 4.2); nâng cấp đổi sprite/prefab theo cấp.")
lead("Config no-code: ", "mỗi cấp (chi phí, hiệu ứng, hình ảnh) khai báo trong ScriptableObject "
     "→ designer tinh chỉnh không cần code.")
para("Danh sách công trình & hiệu ứng nâng cấp:", italic=True, space_after=2)
table(["Công trình", "Hiệu ứng nâng cấp (cấp 1 → N)"], [
    ["Bếp / Trạm nấu", "Nấu nhanh hơn + thêm slot hàng đợi + mở khoá món cao cấp"],
    ["Kho (Storage)", "Tăng sức chứa nguyên liệu & thành phẩm"],
    ["Chuồng trại", "Tăng số vật nuôi & tốc độ sản xuất"],
    ["Vùng trồng (Ruộng)", "Mở thêm ô trồng trên map"],
    ["Quầy bán", "Tăng số đơn cùng lúc & giá bán tốt hơn"],
    ["Nhà chính", "Tăng giới hạn tổng thể & mở tính năng mới"],
], widths=[1.8, 4.5])

# =================================================================
# 5. KIEN TRUC KY THUAT
# =================================================================
h1("5. KIẾN TRÚC KỸ THUẬT & TÍNH KHẢ THI")
para("Phần trọng tâm: dự án không chỉ là ý tưởng mà đã có nền móng mã nguồn vận hành.")

h2("5.1. Nguyên tắc kiến trúc")
lead("Service pattern: ", "mỗi hệ thống lớn là 1 Service, giao tiếp qua interface.")
lead("Event-driven: ", "module phát/nghe Payload events → không phụ thuộc chéo.")
lead("Data-driven: ", "cấu hình nằm trong ScriptableObject → chỉnh không cần code.")

h2("5.2. Vì sao dùng VContainer?")
table(["Tiêu chí", "Tự setup thủ công", "VContainer (đã chọn)"], [
    ["Quản lý vòng đời", "Dễ rò rỉ & sai thứ tự khởi tạo", "LifetimeScope rõ ràng"],
    ["Phụ thuộc", "Ẩn, khó nhìn", "Tường minh qua [Inject]"],
    ["Kiểm thử / mock", "Khó thay implementation", "Dễ swap interface → unit test"],
    ["Hiệu năng", "—", "Cao, ít GC alloc (hợp mobile)"],
], widths=[1.5, 2.4, 2.4])

h2("5.3. Các module nền đã có (Proof of Feasibility)")
table(["Module mã nguồn", "Phục vụ hệ thống", "Trạng thái"], [
    ["Module/Input", "Điều khiển context-aware (4.1)", "Đã chạy — event-driven"],
    ["Module/Map", "Map/Tilemap (4.2), Farm (4.4)", "Đã chạy — grid + preview"],
    ["Module/Time", "Farm (4.4), Cooking (4.5)", "Đã chạy — clock + sync + chống chỉnh giờ"],
    ["Module/UI System", "UI/UX (4.11)", "Đang dựng — theo UICollection"],
    ["Module/DesignPattern", "Hạ tầng chung (MonoSingleton)", "Đã có"],
], widths=[1.7, 3.0, 1.6])

h2("5.4. Map System — Unity Tilemap")
lead("Cấu trúc: ", "Grid + nhiều Tilemap layer (đất, vùng trồng, trang trí, va chạm).")
lead("Trồng lên tile: ", "ô tile hợp lệ ↔ slot trồng; map data lưu trạng thái từng ô.")
lead("Drag & clamp: ", "camera pan theo kéo, giới hạn trong biên map.")
lead("Config no-code: ", "đổi layout/đối tượng qua Tilemap + ScriptableObject.")

h2("5.5. Time System — Chống Cheat & Offline/Online")
lead("Nguồn thời gian: ", "Firebase server timestamp làm mốc tham chiếu (authoritative khi online).")
lead("Chống chỉnh giờ: ", "so sánh delta client–server; ClockService có cờ phát hiện tua giờ.")
lead("Offline: ", "dùng thời gian cache + tích luỹ tiến trình; khi online lại thì sync & "
     "validate để chặn gian lận.")

h2("5.6. Event System — LifeCycle chung, config no-code")
lead("LifeCycle: ", "Schedule → Activate → Progress → Reward → Cleanup (interface chung).")
lead("Config: ", "EventConfig (ScriptableObject) khai báo thời gian, nội dung, thưởng, theme — "
     "thêm event mới không đụng code lõi.")

h2("5.7. UI System — UICollection")
para("UI tổ chức theo Layer (Main / Overlay / Popup) & Group → quản lý thứ tự hiển thị, vòng "
     "đời mở/đóng tập trung, dễ thêm màn hình mới mà không sửa code lõi.")

h2("5.8. Backend & Lưu trữ — Firebase")
lead("Auth + Cloud Save: ", "lưu tiến trình online; offline lưu local rồi đồng bộ khi có mạng.")
lead("Anti-cheat thời gian: ", "tận dụng server timestamp (mục 5.5).")
lead("Tương lai: ", "backend riêng (authoritative time, leaderboard) khi mở rộng.")

h2("5.9. Công nghệ sử dụng (Tech Stack)")
table(["Hạng mục", "Lựa chọn", "Lý do"], [
    ["Engine", "Unity (URP)", "2.5D/2D, build mobile tốt, 60FPS landscape"],
    ["Map", "Unity Tilemap (Grid)", "1 map config được, trồng theo ô, drag"],
    ["DI Container", "VContainer", "Decoupled, testable, ít GC (mục 5.2)"],
    ["Dữ liệu cấu hình", "ScriptableObject", "Recipe/Event/Map config no-code"],
    ["Backend / Save", "Firebase", "Auth + Cloud Save + server time"],
    ["UI", "UICollection (Layer/Group)", "Quản lý màn hình tập trung"],
    ["Source control", "Git (branch theo feature)", "—"],
], widths=[1.5, 2.1, 2.7])

# =================================================================
# 6. NOI DUNG GAME
# =================================================================
h1("6. NỘI DUNG & CÂN BẰNG (CONTENT DESIGN)")
h2("6.1. Cân bằng thời gian & hình thái")
lead("Balance: ", "thời gian trồng/nuôi và độ phức tạp món tăng dần theo tiến trình; "
     "món cao cấp cần nhiều nguyên liệu + bước chế biến hơn nhưng giá bán cao hơn.")
lead("Hình thái sinh trưởng: ", "mỗi cây/vật nuôi có sprite/prefab theo 3 giai đoạn để người "
     "chơi đọc được tiến trình bằng mắt.")
lead("Nhịp 'thích cái mới': ", "khoảng cách giữa các mốc unlock được canh để người chơi luôn "
     "có mục tiêu kế tiếp trong tầm với.")

# =================================================================
# 7. MVP SCOPE
# =================================================================
h1("7. PHẠM VI MVP (SCOPE CONTROL)")
para("TRONG SCOPE — Demo dọc 8 tuần:", bold=True, space_after=2)
bullet("Map System (Tilemap, config, trồng lên tile, drag map).")
bullet("Core loop: chọn món → checklist → trồng/nuôi → thu hoạch → nấu (animation) → cutscene.")
bullet("Unlock System (nguyên liệu + món, mở bằng coins) + Quest ép làm món vừa unlock.")
bullet("Update System: xây & nâng cấp công trình trong farm (Bếp, Kho, Chuồng, Ruộng, Quầy bán).")
bullet("Tutorial Bánh Mì + 4–6 món Việt + 3–5 nguyên liệu/ vật nuôi.")
bullet("Event System (lifecycle + config) với 1 event 'Tết: Bánh Chưng'.")
bullet("Cooking animation, Shop booster, Time System chống cheat, Save offline/online (Firebase).")
para("NGOÀI SCOPE — hướng tương lai:", bold=True, space_after=2)
bullet("Đa map/vùng miền, nhiều event (Trung Thu…), mini-game lễ hội.")
bullet("IAP/monetization, leaderboard, backend riêng (authoritative time/anti-cheat server).")
bullet("Đa ngôn ngữ.")

# =================================================================
# 8. QUAN TRI DU AN
# =================================================================
h1("8. QUẢN TRỊ DỰ ÁN")
h2("8.1. Đối tượng & Thị trường")
para("Game thủ casual mobile thích farming/cooking (16–35), ưa thẩm mỹ indie/hand-drawn. "
     "Phân khúc cosy farming/cooking đang tăng. Lợi thế: ngách ẩm thực Việt Nam chưa đối thủ chiếm.")

h2("8.2. Tiến độ — Scrum, 4 Sprint × 2 tuần (từ 02/06/2026)")
table(["Sprint", "Thời gian", "Sprint Goal", "Backlog chính"], [
    ["Sprint 1", "02–15/06", "Map + Core Loop nền", "Tilemap config + drag map; di chuyển/one-tap; Time System nền"],
    ["Sprint 2", "16–29/06", "Farm + Unlock", "Trồng trọt/chăn nuôi theo Time + giai đoạn sinh trưởng; Unlock nguyên liệu"],
    ["Sprint 3", "30/06–13/07", "Cooking + Quest + Update + Shop", "Cooking animation; Unlock món; Quest; Tutorial bánh mì; Update System (nâng cấp công trình); Shop booster"],
    ["Sprint 4", "14–27/07", "Event + Cutscene + Online + Polish", "Event Tết bánh chưng; cutscene; Firebase save offline/online; UI; build demo"],
], widths=[0.8, 1.0, 1.7, 2.8])

h2("8.3. Tiêu chí Demo Thành công (KPI)")
bullet("Hoàn thành trọn 1 vòng core loop (chọn món → trồng/nuôi → nấu → cutscene → coins) không lỗi chặn.")
bullet("Unlock + Quest hoạt động: mở món mới bằng coins, quest ép làm món đúng số lần.")
bullet("≥ 4 món Việt làm được; nguyên liệu có 3 giai đoạn sinh trưởng hiển thị đúng.")
bullet("Event 'Tết: Bánh Chưng' kích hoạt & nhận thưởng được; config qua ScriptableObject.")
bullet("Update System: nâng cấp ≥ 1 công trình (vd Bếp/Kho) → buff áp dụng đúng (nấu nhanh/kho lớn hơn).")
bullet("Tutorial bánh mì dẫn người mới qua core loop < 3 phút.")
bullet("Time System chống cheat hoạt động; Save offline/online (Firebase) giữ tiến trình.")
bullet("Chạy ổn định 60 FPS, màn hình ngang, trên mobile tầm trung.")

h2("8.4. Rủi ro & Phương án Dự phòng")
table(["Rủi ro", "Mức", "Phương án"], [
    ["Quỹ thời gian hạn chế (đi làm song song)", "Cao", "Break task nhỏ; ưu tiên KPI theo Sprint"],
    ["Art hand-drawn nhiều món/giai đoạn", "Cao", "3 SV Đa phương tiện song song; placeholder + asset free giai đoạn đầu"],
    ["Phình phạm vi (scope creep)", "Cao", "Bám mục 7; event/đa map → backlog tương lai"],
    ["Cân bằng số liệu (time/giá)", "TB", "Số liệu trong ScriptableObject → tinh chỉnh nóng"],
    ["Đồng bộ offline/online & chống cheat", "TB", "Bọc qua service; validate khi reconnect; tận dụng Firebase time"],
    ["Bản quyền nhân vật cutscene", "TB", "Dùng nhân vật gốc/được cấp phép"],
], widths=[2.4, 0.7, 3.2])

# =================================================================
# 9. CHECKLIST PHAN CONG
# =================================================================
h1("9. CHECKLIST PHÂN CÔNG (GDD / DEV / ART)")
h2("9.1. Checklist — Game Design")
table(["Hạng mục", "Mô tả", "Mục"], [
    ["Hệ thống Unlock", "Nguyên liệu + món; checklist nguyên liệu; unlock = coins", "4.3"],
    ["Core Loop", "List nguyên liệu → trồng/nuôi → thu hoạch → sản xuất → cutscene", "3"],
    ["Hệ thống Quest", "Ép user làm món vừa unlock", "4.6"],
    ["Hệ thống Event", "Làm Tết bánh chưng trước (quảng bá văn hoá)", "4.7"],
    ["Hệ thống Tutorial", "Dễ hiểu + món thân thuộc (bánh mì)", "4.8"],
    ["Cutscene sáng tạo", "Đồ ăn gắn nhân vật nổi tiếng", "4.10"],
    ["Balance", "Cân bằng thời gian + hình thái trồng trọt/sản xuất", "6.1"],
], widths=[1.8, 3.7, 0.8])
h2("9.2. Checklist — Developer")
table(["Hệ thống", "Yêu cầu", "Mục"], [
    ["Map System", "1 map, config được, trồng lên map, drag map; theo Tilemap Unity", "4.2 / 5.4"],
    ["Farm System", "Trồng trọt/chăn nuôi theo Time System", "4.4"],
    ["Cooking System", "Animation cho cooking", "4.5"],
    ["Time System", "Chống cheat — Firebase time, sync timing", "5.5"],
    ["Shop", "Booster (x2 time trồng, x2 resource…)", "4.9"],
    ["Event System", "LifeCycle chung, config dễ (không đụng code)", "5.6"],
    ["Update System", "Xây & nâng cấp công trình trong farm (đổi hình theo cấp)", "4.12"],
], widths=[1.5, 4.0, 0.8])
h2("9.3. Checklist — Art")
table(["Hạng mục", "Mô tả"], [
    ["Tile cho map", "Bộ tile nền cho Tilemap"],
    ["Environment", "Design map → thêm components trang trí"],
    ["Nguyên liệu", "Cây trồng & vật nuôi theo Progress (giống / phát triển / trưởng thành)"],
    ["Công trình", "Bếp/Kho/Chuồng/Quầy… với hình theo từng cấp nâng cấp (cấp 1 → N)"],
    ["Cutscene", "Thành phẩm → câu chuyện"],
    ["UI/UX", "Bộ UI màn hình ngang, icon món/nguyên liệu"],
], widths=[1.6, 4.7])

# =================================================================
# 10. POST-MVP
# =================================================================
h1("10. HƯỚNG MỞ RỘNG (POST-MVP)")
bullet("Thêm món & nguyên liệu vùng miền; đa map/khu vực ẩm thực Bắc–Trung–Nam.")
bullet("Đa event: Trung Thu (bánh trung thu), lễ hội ẩm thực; mini-game nấu ăn (rhythm).")
bullet("Kinh tế động (chợ giá cung–cầu), thời tiết/mùa vụ.")
bullet("Monetization (cosmetic, booster), leaderboard, backend riêng (authoritative time/anti-cheat).")
bullet("Bộ sưu tập cutscene 'Câu chuyện món Việt' — quảng bá văn hoá.")

para("", space_after=0)
para("— HẾT —", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Tài liệu cập nhật ngày 18/06/2026 — Vũ Đình Khoa (PO / Lead Dev / Game Design).",
     italic=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(DST)
print("OK saved:", DST)
