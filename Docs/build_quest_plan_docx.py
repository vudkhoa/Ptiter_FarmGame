from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Docs/Quest_Module_Research_Plan.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_run_font(run, name="Calibri", size=11, color=None, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_body(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if level == 1:
        size, color, before, after = 16, "2E74B5", 18, 10
    elif level == 2:
        size, color, before, after = 13, "2E74B5", 14, 7
    else:
        size, color, before, after = 12, "1F4D78", 10, 5
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    for run in p.runs:
        set_run_font(run, size=size, color=color, bold=True)
    return p


def add_callout(doc, title, body, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, color="1F3A5F", bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    r2 = p2.add_run(body)
    set_run_font(r2)
    doc.add_paragraph()


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], "E8EEF5")
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=10, color="0B2545", bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if idx == 0 and len(headers) > 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            set_run_font(r, size=10)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    return doc


def build():
    doc = setup_document()

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("Quest Module Research & Implementation Plan")
    set_run_font(title_run, size=22, color="0B2545", bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("Ptiter FarmGame - research proposal for a data-driven, testable Quest module")
    set_run_font(r, size=11, color="555555", italic=True)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    rows = [
        ("Status", "Research / proposal only. No implementation yet."),
        ("Primary goal", "Design a Quest module aligned with current module style: MessagePipe events, VContainer DI, ScriptableObject data, and Unity Test Runner coverage."),
        ("Scope now", "No UI/view. Data-driven service logic, RAM-backed progress storage, adapters, payloads, and unit tests."),
        ("Recommended preset", "compact_reference_guide: dense but readable technical review document."),
    ]
    for idx, row in enumerate(rows):
        for col, value in enumerate(row):
            cell = meta.cell(idx, col)
            if col == 0:
                set_cell_shading(cell, "E8EEF5")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=10, bold=(col == 0), color="0B2545" if col == 0 else None)
    set_table_width(meta, [1.55, 4.95])

    add_callout(
        doc,
        "Recommendation",
        "Build Quest as a reusable domain module. Keep Farm-specific interpretation in adapters, not in Quest core. QuestService should consume generic quest events and objective handlers should be resolved through a small registry injected by VContainer.",
    )

    add_heading(doc, "1. Research Findings", 1)
    add_body(
        doc,
        "Farming and life-sim games commonly use quests to teach the production loop, create short-term goals, drive long-term progression, and give designers a tunable layer above crops, animals, inventory, and time.",
    )
    add_matrix(
        doc,
        ["Quest type", "Typical use in farming games", "MVP relevance"],
        [
            ("Tutorial / Story", "Teach planting, harvesting, feeding, storage, and first production loops.", "High"),
            ("Request Board", "Short tasks such as deliver item, gather material, or produce item within a small time window.", "Medium"),
            ("Special Order", "Larger weekly-style goals, often requiring progress after accepting the quest.", "Medium"),
            ("Farm Production", "Harvest X crops, feed X animals, collect X products, or produce specific goods.", "High"),
            ("Progression / Unlock", "Reach milestones that unlock areas, recipes, buildings, or systems.", "Medium"),
            ("Achievement-style", "Long-term cumulative targets that reward mastery and retention.", "Low for MVP"),
            ("Delivery / Turn-in", "Have or consume items from inventory to satisfy an NPC/order.", "High"),
        ],
        [1.45, 3.35, 1.7],
    )

    add_heading(doc, "2. Design Principles", 1)
    principles = [
        "Quest core must be decoupled from Farm so the module can be reused in another game.",
        "Farm integration should live in a Quest-Farm adapter layer that converts Farm events into generic Quest events.",
        "Use MessagePipe for state changes and domain events. Add payloads only when listeners need context.",
        "Use VContainer for instances and registries. Avoid static singletons and direct new calls inside production logic.",
        "Separate ScriptableObject data from runtime state and objective evaluation logic.",
        "Start with RAM-backed storage, but put it behind an interface so persistent storage can replace it later.",
        "Unit tests should validate behavior through service APIs and MessagePipe-like test publishers/subscribers.",
    ]
    for item in principles:
        add_bullet(doc, item)

    add_heading(doc, "3. Proposed Module Structure", 1)
    add_body(doc, "Suggested folder layout:")
    add_matrix(
        doc,
        ["Folder", "Responsibility"],
        [
            ("Scripts/SO", "QuestCatalogSO, QuestDefinitionSO, reward/objective data assets."),
            ("Scripts/Runtime", "Runtime state records: active quest, objective progress, completion flags."),
            ("Scripts/Service", "IQuestService and QuestService orchestration."),
            ("Scripts/Logic", "Objective handlers and QuestObjectiveHandlerRegistry."),
            ("Scripts/Storage", "IQuestProgressStorage and InMemoryQuestProgressStorage."),
            ("Scripts/Payloads", "QuestAccepted, QuestProgressChanged, QuestCompleted, QuestRewardClaimed, QuestExpired."),
            ("Scripts/Adapters", "Farm-to-Quest and Storage-to-Quest event adapters."),
            ("Tests", "Unity Test Runner edit-mode tests for Quest FSM and adapters."),
        ],
        [1.9, 4.6],
    )

    add_heading(doc, "4. ScriptableObject Data Model", 1)
    add_body(
        doc,
        "Quest data should be authored by designers as ScriptableObject assets. Runtime progress should not be stored in SO files; SO files are definitions only.",
    )
    add_matrix(
        doc,
        ["SO / data class", "Important fields"],
        [
            ("QuestCatalogSO", "List<QuestDefinitionSO> quests"),
            ("QuestDefinitionSO", "questId, title, description, category, repeatPolicy, activationType, timeLimitSeconds, prerequisites, objectives, rewards"),
            ("QuestObjectiveData", "objectiveId, objectiveType, targetId, requiredAmount, progressMode, optional tags"),
            ("QuestRewardData", "coins, itemRewards, unlockIds, optional followUpQuestIds"),
            ("QuestPrerequisiteData", "required quest ids, min level/rank, feature flags, optional custom conditions"),
        ],
        [2.0, 4.5],
    )

    add_heading(doc, "5. Quest Categories & Objective Types", 1)
    add_matrix(
        doc,
        ["Category", "Objective examples", "Notes"],
        [
            ("Farm Crop", "PlantCrop, HarvestCrop", "Should listen to semantic farm events, not raw slot mutation."),
            ("Farm Animal", "BuyAnimal, FeedAnimal, CollectAnimalProduct", "Animal lifecycle differs from crop lifecycle."),
            ("Inventory", "GainItem, ReachInventoryAmount, TurnInItem", "Can use InventoryChangedPayload and IStorageService."),
            ("Economy", "EarnCoins, SpendCoins, ReachCoins", "Current storage has Coins but no coin-changed payload yet."),
            ("Progression", "CompleteQuest, UnlockFeature, ReachRank", "Useful after MVP."),
            ("Timed", "CompleteWithinSeconds, DailyReset, WeeklyOrder", "Needs clock/calendar policy later."),
        ],
        [1.45, 2.55, 2.5],
    )
    add_body(
        doc,
        "Recommended MVP objective types: PlantCrop, HarvestCrop, FeedAnimal, CollectAnimalProduct, GainItem, ReachInventoryAmount, TurnInItem, and EarnCoins/SpendCoins only if a coin event is added.",
    )

    add_heading(doc, "6. Runtime Service Design", 1)
    add_body(doc, "QuestService should be the only public write path for quest state. It accepts generic quest events and delegates objective-specific progress rules to handlers.")
    for item in [
        "IQuestService.AcceptQuest(questId) creates runtime state through IQuestProgressStorage.",
        "IQuestService.ReportEvent(QuestEvent evt) updates only active quests.",
        "QuestObjectiveHandlerRegistry resolves a handler by objective type.",
        "QuestService publishes progress/completion/reward events through MessagePipe.",
        "Rewards are claimed separately from completion to prevent accidental double grants.",
    ]:
        add_number(doc, item)

    add_heading(doc, "7. Event & Payload Plan", 1)
    add_matrix(
        doc,
        ["Payload", "When published", "Key data"],
        [
            ("QuestAcceptedPayload", "Quest enters active state.", "questId"),
            ("QuestProgressChangedPayload", "Any objective progress changes.", "questId, objectiveId, current, required"),
            ("QuestCompletedPayload", "All objectives are complete.", "questId"),
            ("QuestRewardClaimedPayload", "Rewards are granted.", "questId, coins, item rewards"),
            ("QuestExpiredPayload", "Timed quest exceeds deadline.", "questId, expiredAtUtcTicks"),
            ("QuestEventPayload", "Optional generic event bus input.", "eventType, targetId, amount, source, timestamp"),
        ],
        [1.9, 2.25, 2.35],
    )

    add_heading(doc, "8. Farm Integration Without Tight Coupling", 1)
    add_callout(
        doc,
        "Current risk",
        "FarmSlotChangedPayload is too low-level for Quest. After harvest, crop entityId may be reset, so Quest cannot reliably know what was harvested from slot state alone.",
        fill="FFF2CC",
    )
    add_body(doc, "Recommended Farm semantic events:")
    for item in [
        "FarmEntityPlantedPayload(cell, entityId, isAnimal)",
        "FarmEntityFedPayload(cell, animalId, consumedItemId)",
        "FarmEntityHarvestedPayload(cell, entityId, isAnimal, productItemId, amount)",
    ]:
        add_bullet(doc, item)
    add_body(
        doc,
        "QuestFarmAdapter subscribes to those Farm events, converts them into generic QuestEvent instances, then calls IQuestService.ReportEvent(). Quest core stays ignorant of Farm-specific classes.",
    )

    add_heading(doc, "9. Storage Strategy", 1)
    add_body(
        doc,
        "The existing Storage module currently exposes IStorageService for coins, inventory, cheat state, and Save(). It does not yet expose quest progress. For this phase, keep quest state in RAM behind a Quest storage interface.",
    )
    add_matrix(
        doc,
        ["Option", "Pros", "Cons", "Recommendation"],
        [
            ("Quest-owned RAM storage", "Minimal changes; easy to test; replaceable later.", "Not persisted after restart.", "Use for MVP"),
            ("Extend Storage module", "Centralized persistence contract.", "Touches shared module before persistence design is stable.", "Defer"),
            ("Embed in PlayerData now", "Closer to final save path.", "Higher blast radius; more migration concerns.", "Defer"),
        ],
        [1.35, 1.85, 1.75, 1.55],
    )

    add_heading(doc, "10. Unit Test Plan", 1)
    tests = [
        "Accept quest creates active runtime state and publishes QuestAcceptedPayload.",
        "Events do not progress inactive quests.",
        "PlantCrop progresses only matching targetId.",
        "HarvestCrop progresses only when event source is harvest, not inventory gain alone.",
        "Multi-objective quest completes only when every objective is complete.",
        "GainItem reacts to InventoryChangedPayload through a Storage adapter.",
        "TurnInItem consumes inventory through IStorageService.RemoveItem.",
        "ClaimReward grants coins/items once and publishes QuestRewardClaimedPayload.",
        "InMemoryQuestProgressStorage can restore active/completed states in the same test session.",
        "Farm adapter converts planted/fed/harvested events into generic QuestEvent correctly.",
    ]
    for item in tests:
        add_bullet(doc, item)

    add_heading(doc, "11. Phased Roadmap", 1)
    add_matrix(
        doc,
        ["Phase", "Deliverables", "Exit criteria"],
        [
            ("Research approval", "Review this plan and confirm MVP objective types.", "No open architecture disagreement."),
            ("Core MVP", "Quest SOs, QuestService, RAM storage, handlers, payloads.", "All core tests pass."),
            ("Farm adapter", "Semantic farm events and QuestFarmAdapter.", "Farm event tests pass without Quest referencing Farm internals."),
            ("Storage evolution", "Persistent quest state contract or PlayerData integration.", "Quest state survives reload."),
            ("UI later", "Quest journal, board, progress notifications.", "UI listens to payloads; no logic in view."),
        ],
        [1.25, 3.0, 2.25],
    )

    add_heading(doc, "12. Open Questions For Review", 1)
    for item in [
        "Should MVP include timed quests now, or keep deadline fields dormant until calendar/time design is clearer?",
        "Should Quest module own IQuestProgressStorage, or should Storage module define the storage contract?",
        "Should FarmService emit semantic events as part of this task, or should the first Quest tests use fake adapter events?",
        "Do designers need random/generated quests now, or only authored QuestDefinitionSO assets?",
        "Should quest rewards immediately grant on completion, or require explicit claim?",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "13. Research Sources", 1)
    sources = [
        ("Stardew Valley Wiki - Quests", "https://stardewvalleywiki.com/Quests"),
        ("Stardew Valley Wiki - Special Orders", "https://stardewvalleywiki.com/Quests#List_of_Special_Orders"),
        ("Story of Seasons common gameplay elements", "https://en.wikipedia.org/wiki/Story_of_Seasons"),
        ("FarmVille gameplay loop", "https://en.wikipedia.org/wiki/FarmVille"),
        ("Fields of Mistria gameplay and request board summary", "https://en.wikipedia.org/wiki/Fields_of_Mistria"),
    ]
    for label, url in sources:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        add_hyperlink(p, label, url)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Quest Module Research Plan - Ptiter FarmGame")
    set_run_font(fr, size=9, color="555555")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
